import os
from docx import Document

docx_dir = r"d:\audi dev\docx"
with open("docx_output.txt", "w", encoding="utf-8") as f:
    for filename in os.listdir(docx_dir):
        if filename.endswith(".docx"):
            f.write(f"\n{'='*50}\nReading: {filename}\n{'='*50}\n")
            try:
                doc = Document(os.path.join(docx_dir, filename))
                for i, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text:
                        f.write(text + "\n")
                    if i > 200:
                        f.write(f"... (truncated {len(doc.paragraphs)-i} more paragraphs)\n")
                        break
            except Exception as e:
                f.write(f"Error reading {filename}: {e}\n")
